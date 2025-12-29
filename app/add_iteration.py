# add_iteration.py
# 添加新的迭代记录到Word文档

from docx import Document
from datetime import datetime
import os
import sys

def add_iteration(version, changes, file_structure=None, tech_details=None, impact=None, test_notes=None, remarks=None):
    """
    添加新的迭代记录
    
    参数:
    version: 版本号 (如 "1.1")
    changes: 变更内容描述
    file_structure: 文件结构变化 (可选)
    tech_details: 技术细节列表 (可选)
    impact: 影响范围 (可选)
    test_notes: 测试说明 (可选)
    remarks: 备注 (可选)
    """
    doc_path = '代码迭代记录.docx'
    
    # 如果文档不存在，先创建基础文档
    if not os.path.exists(doc_path):
        print(f"错误：找不到 {doc_path}，请先运行 create_iteration_log.py 创建文档")
        return False
    
    # 打开现有文档
    doc = Document(doc_path)
    
    # 添加分隔线
    doc.add_paragraph('=' * 50)
    
    # 添加新迭代标题
    doc.add_heading(f'迭代 {version}', level=1)
    
    # 版本信息
    p_info = doc.add_paragraph()
    p_info.add_run('版本号：').bold = True
    p_info.add_run(version)
    p_info.add_run(' | ')
    p_info.add_run('日期：').bold = True
    p_info.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    p_info.add_run(' | ')
    p_info.add_run('开发者：').bold = True
    p_info.add_run('Auto')
    
    # 变更内容
    doc.add_heading('变更内容', level=2)
    doc.add_paragraph(changes)
    
    # 文件结构变化
    if file_structure:
        doc.add_heading('文件结构变化', level=2)
        if isinstance(file_structure, str):
            doc.add_paragraph(file_structure)
        elif isinstance(file_structure, list):
            for item in file_structure:
                doc.add_paragraph(item, style='List Bullet')
    
    # 技术细节
    if tech_details:
        doc.add_heading('技术细节', level=2)
        if isinstance(tech_details, str):
            doc.add_paragraph(tech_details)
        elif isinstance(tech_details, list):
            for item in tech_details:
                doc.add_paragraph(item, style='List Bullet')
    
    # 影响范围
    if impact:
        doc.add_heading('影响范围', level=2)
        if isinstance(impact, str):
            doc.add_paragraph(impact)
        elif isinstance(impact, list):
            for item in impact:
                doc.add_paragraph(item, style='List Bullet')
    
    # 测试说明
    if test_notes:
        doc.add_heading('测试说明', level=2)
        if isinstance(test_notes, str):
            doc.add_paragraph(test_notes)
        elif isinstance(test_notes, list):
            for item in test_notes:
                doc.add_paragraph(item, style='List Bullet')
    
    # 备注
    if remarks:
        doc.add_heading('备注', level=2)
        doc.add_paragraph(remarks)
    
    # 保存文档
    doc.save(doc_path)
    print(f'迭代记录 {version} 已添加到文档')
    return True

if __name__ == "__main__":
    # 示例用法
    if len(sys.argv) < 3:
        print("用法: python add_iteration.py <版本号> <变更内容>")
        print("示例: python add_iteration.py 1.1 '修复了文档管理的bug'")
        sys.exit(1)
    
    version = sys.argv[1]
    changes = sys.argv[2]
    
    add_iteration(version, changes)

